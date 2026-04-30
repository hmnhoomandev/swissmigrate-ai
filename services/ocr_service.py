from io import BytesIO
from pathlib import Path


SUPPORTED_EXTENSIONS = (".pdf", ".png", ".jpg", ".jpeg", ".webp", ".docx")
MAX_OCR_PDF_PAGES = 8


def _ocr_image(image) -> str:
    import pytesseract

    if pytesseract.pytesseract.tesseract_cmd in {"", "tesseract"}:
        executable = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
        if executable.exists():
            pytesseract.pytesseract.tesseract_cmd = str(executable)
    return pytesseract.image_to_string(image).strip()


def _ocr_pdf_pages(payload: bytes) -> str:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(payload)
    page_count = min(len(pdf), MAX_OCR_PDF_PAGES)
    page_text = []
    for page_index in range(page_count):
        page = pdf[page_index]
        bitmap = page.render(scale=2.4).to_pil()
        text = _ocr_image(bitmap)
        if text:
            page_text.append(text)
    return "\n\n".join(page_text).strip()


def validate_upload(uploaded_file) -> list[str]:
    if uploaded_file is None:
        return []

    warnings = []
    name = uploaded_file.name.lower()

    if not name.endswith(SUPPORTED_EXTENSIONS):
        warnings.append("Unsupported file type. Please upload a PDF, DOCX, JPG, JPEG, PNG, or WEBP file.")
    return warnings


def extract_text_from_upload(uploaded_file) -> tuple[str, list[str]]:
    if uploaded_file is None:
        return "", []

    warnings = validate_upload(uploaded_file)
    if warnings:
        return "", warnings

    name = uploaded_file.name.lower()
    payload = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        try:
            import pdfplumber

            with pdfplumber.open(BytesIO(payload)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n\n".join(page.strip() for page in pages if page.strip())
            if text:
                return text, warnings
            warnings.append("No selectable text found in the PDF. Trying OCR on the scanned pages.")
        except Exception as exc:
            warnings.append(f"PDF extraction failed: {exc}")

        try:
            text = _ocr_pdf_pages(payload)
            if text:
                return text, warnings
            warnings.append("OCR could not find readable text in the PDF.")
        except Exception as exc:
            warnings.append(
                "PDF OCR is unavailable. Install the Tesseract system package and make sure it is in PATH. "
                f"Details: {exc}"
            )
        return "", warnings

    if name.endswith(".docx"):
        try:
            from docx import Document

            document = Document(BytesIO(payload))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            tables = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables.append(" | ".join(cells))
            text = "\n".join(paragraphs + tables).strip()
            if text:
                return text, warnings
            warnings.append("No readable text was found in the Word document.")
        except Exception as exc:
            warnings.append(f"Word document extraction failed: {exc}")
        return "", warnings

    try:
        from PIL import Image

        image = Image.open(BytesIO(payload))
        text = _ocr_image(image)
        if text:
            return text, warnings
        warnings.append("No readable text was found in the image.")
        return "", warnings
    except Exception as exc:
        warnings.append(
            "Image OCR is unavailable. Install Pillow, pytesseract, and the Tesseract system package. "
            f"Details: {exc}"
        )
        return "", warnings
